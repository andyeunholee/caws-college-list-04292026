"""Magazine-style Word (.docx) report assembler — supports English and Korean.

Layout:
- Page header (student name) + "Page X of Y" footer
- Title block + Student Profile Snapshot + Note on Probabilities
- Part 1  — Home-state colleges (LACs included, tagged "(LAC)")
- Part 2  — National universities (excluding home state, excluding LACs)
- Part 3  — Liberal Arts Colleges nationwide (excluding home-state LACs)
- Part 4  — Early Decision options
- Part 5  — Early Action options
- Part 6  — Action plan
- Appendix — Glossary
"""

from __future__ import annotations

import datetime as _dt
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from .i18n import Lang, s
from .schema import CollegeRow, Scope, StudentProfile, TieredList, ValidationFlag
from .utils.glossary import get_glossary


# ────────────────────────── color & style constants ──────────────────────────
HEADER_BG = "1F4E79"
HEADER_TXT = "FFFFFF"
ACCENT = "1F4E79"
ZEBRA_BG = "F4F6F8"
NOTE_BG = "FFF8E1"
WARN_TXT = "8A6D3B"
SUBTITLE_GRAY = "555555"

TIER_COLOR = {
    "reach": "C00000",
    "match": "1F6FB2",
    "safety": "2E7D32",
}
TIER_ICON = {
    "reach": "🎯",
    "match": "⚖",
    "safety": "🛡",
}

ED_MAX = 50
EA_MAX = 50

# QuestBridge National College Match partner colleges (AY 2025-26).
QB_PARTNERS: frozenset[str] = frozenset({
    "Amherst College", "Barnard College", "Bates College", "Bowdoin College",
    "Brown University", "Bryn Mawr College", "California Institute of Technology",
    "Carleton College", "Claremont McKenna College", "Colby College",
    "Colgate University", "Colorado College", "Columbia University",
    "Cornell University", "Dartmouth College", "Davidson College",
    "Denison University", "Duke University", "Emory University", "Grinnell College",
    "Hamilton College", "Haverford College", "Johns Hopkins University",
    "Kenyon College", "Macalester College", "Massachusetts Institute of Technology",
    "Middlebury College", "Northwestern University", "Oberlin College",
    "Pomona College", "Princeton University", "Rice University", "Scripps College",
    "Skidmore College", "Smith College", "Stanford University", "Swarthmore College",
    "College of the Holy Cross", "Trinity College", "Tufts University",
    "University of Chicago", "University of Notre Dame", "University of Pennsylvania",
    "University of Southern California", "University of Virginia",
    "University of North Carolina at Chapel Hill", "Vanderbilt University",
    "Vassar College", "Washington and Lee University", "Washington & Lee University",
    "Washington University in St. Louis", "Wellesley College", "Wesleyan University",
    "Williams College", "Yale University", "Carnegie Mellon University",
    "Boston College",
})


# Korean UI font. 맑은 고딕 (Malgun Gothic) ships with Windows and Office for
# Mac; Word substitutes a sans-serif Hangul font where it is missing.
KO_FONT = "맑은 고딕"


def _set_east_asian_font(rPr, font_name: str) -> None:
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _is_qb(name: str) -> bool:
    return name in QB_PARTNERS


def _tier_label(lang: Lang, tier: str) -> str:
    return s(lang, f"tier_{tier}")


# ────────────────────────── primitive XML helpers ──────────────────────────


def _shade_cell(cell, hex_color: str, text_white: bool = False) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)
    if text_white:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor.from_string(HEADER_TXT)


def _set_cell_text(
    cell,
    text: str,
    bold: bool = False,
    size: int = 10,
    color_hex: str | None = None,
    align: int | None = None,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_bottom_border(paragraph, color_hex: str = HEADER_BG, sz: str = "8") -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_field(paragraph, instr: str, placeholder: str = "1"):
    """Insert a complete Word field (e.g., PAGE, NUMPAGES) into a paragraph.

    Emits the full begin / instrText / separate / result / end sequence and
    marks the field dirty so Word recalculates it on open. Without the
    separate+result part, viewers show an empty result (e.g. "Page 5 of ")
    until the user manually updates fields.

    Returns the run holding the visible result so callers can style it.
    """
    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    fldBegin.set(qn("w:dirty"), "true")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f" {instr} "
    fldSep = OxmlElement("w:fldChar")
    fldSep.set(qn("w:fldCharType"), "separate")
    run._r.append(fldBegin)
    run._r.append(instrText)
    run._r.append(fldSep)

    result_run = paragraph.add_run(placeholder)

    end_run = paragraph.add_run()
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    end_run._r.append(fldEnd)
    return result_run


def _usable_width(doc: Document):
    section = doc.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def _apply_column_ratios(table, ratios: list[float], total_width) -> None:
    table.autofit = False
    table.allow_autofit = False
    for col_idx, ratio in enumerate(ratios):
        width = int(total_width * ratio)
        try:
            table.columns[col_idx].width = width
        except IndexError:
            continue
        for row in table.rows:
            row.cells[col_idx].width = width


# ────────────────────────── page-level scaffolding ──────────────────────────


def _setup_document(doc: Document, profile: StudentProfile, lang: Lang) -> None:
    normal = doc.styles["Normal"]
    # Latin text: 맑은 고딕 for KO docs, Calibri for EN docs. Hangul glyphs
    # are drawn with the *East Asian* font slot, which python-docx does not
    # set — without it Word falls back to 바탕 (serif). Set it explicitly.
    normal.font.name = KO_FONT if lang == "ko" else "Calibri"
    normal.font.size = Pt(10.5)
    _set_east_asian_font(normal.element.get_or_add_rPr(), KO_FONT)
    # Also set the document-wide default so header/footer/table text inherits it.
    styles_el = doc.styles.element
    rpr_default = styles_el.find(qn("w:docDefaults") + "/" + qn("w:rPrDefault") + "/" + qn("w:rPr"))
    if rpr_default is not None:
        _set_east_asian_font(rpr_default, KO_FONT)

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run(s(lang, "page_header_template", name=profile.name))
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(SUBTITLE_GRAY)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pre = footer_p.add_run(s(lang, "page_footer_page_prefix"))
    pre.font.size = Pt(9)
    pre.font.color.rgb = RGBColor.from_string(SUBTITLE_GRAY)
    _add_field(footer_p, "PAGE")
    of_run = footer_p.add_run(s(lang, "page_footer_of"))
    of_run.font.size = Pt(9)
    of_run.font.color.rgb = RGBColor.from_string(SUBTITLE_GRAY)
    _add_field(footer_p, "NUMPAGES")
    for r in footer_p.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(SUBTITLE_GRAY)


def _h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(ACCENT)
    _add_bottom_border(p, ACCENT, sz="8")
    p.paragraph_format.space_after = Pt(8)


def _h2(doc: Document, text: str, color_hex: str = ACCENT) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _h3(doc: Document, text: str, color_hex: str = ACCENT) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def _para(doc: Document, text: str, italic: bool = False, size: float = 10.5) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)


def _page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ────────────────────────── content sections ──────────────────────────


def _add_title_block(doc: Document, profile: StudentProfile, lang: Lang) -> None:
    title = doc.add_paragraph()
    run = title.add_run(s(lang, "report_title"))
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(ACCENT)
    _add_bottom_border(title, ACCENT, sz="12")
    title.paragraph_format.space_after = Pt(4)

    sub_bits = [profile.name]
    if profile.high_school:
        sub_bits.append(profile.high_school)
    track = profile.career_goal or profile.intended_major
    if track:
        sub_bits.append(track)
    sub = doc.add_paragraph()
    sub_run = sub.add_run(" · ".join(sub_bits))
    sub_run.italic = True
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor.from_string(SUBTITLE_GRAY)
    sub.paragraph_format.space_after = Pt(2)

    tagline_text = s(
        lang,
        "subtitle_qb_track" if profile.questbridge_candidate else "subtitle_default",
    )
    tagline = doc.add_paragraph()
    tag_run = tagline.add_run(tagline_text)
    tag_run.font.size = Pt(11)
    tag_run.font.color.rgb = RGBColor.from_string(SUBTITLE_GRAY)
    tagline.paragraph_format.space_after = Pt(14)


def _profile_rows(p: StudentProfile, lang: Lang) -> list[tuple[str, str]]:
    coursework_count = len(p.coursework) if p.coursework else None
    leadership_count = len(p.leadership) if p.leadership else (
        len([a for a in p.activities if a.role and "leader" in a.role.lower()]) or None
    )
    award_summary = ", ".join(a.name for a in p.awards[:4]) if p.awards else None
    intended_track = p.career_goal or p.intended_major

    sat_str: str | None = None
    if p.sat_total:
        sat_str = s(
            lang, "profile_label_sat_value",
            total=p.sat_total, ebrw=p.sat_ebrw or "?", math=p.sat_math or "?",
        )

    hs_str: str | None = None
    if p.high_school:
        bits = [p.high_school]
        if p.city:
            bits.append(p.city)
        if p.state:
            bits.append(p.state)
        hs_str = ", ".join(bits)

    rows: list[tuple[str, str | None]] = [
        (s(lang, "profile_label_high_school"), hs_str),
        (s(lang, "profile_label_intended_track"), intended_track),
        (s(lang, "profile_label_uw_gpa"), f"{p.gpa_unweighted}" if p.gpa_unweighted is not None else None),
        (s(lang, "profile_label_w_gpa"), f"{p.gpa_weighted}" if p.gpa_weighted is not None else None),
        (s(lang, "profile_label_class_rank"), p.class_rank),
        (s(lang, "profile_label_sat"), sat_str),
        (s(lang, "profile_label_act"), f"{p.act_composite}" if p.act_composite else None),
        (
            s(lang, "profile_label_ap_ib_de"),
            s(lang, "profile_courses_unit", n=coursework_count) if coursework_count else None,
        ),
        (s(lang, "profile_label_awards"), award_summary),
        (
            s(lang, "profile_label_service"),
            s(lang, "profile_hours_unit", n=p.community_service_hours) if p.community_service_hours else None,
        ),
        (
            s(lang, "profile_label_leadership"),
            s(lang, "profile_positions_unit", n=leadership_count) if leadership_count else None,
        ),
        (
            s(lang, "profile_label_program"),
            s(lang, "profile_value_questbridge_planned") if p.questbridge_candidate else None,
        ),
    ]
    return [(k, v) for k, v in rows if v]


def _add_profile_snapshot(doc: Document, p: StudentProfile, lang: Lang) -> None:
    _h1(doc, s(lang, "profile_h1"))

    rows = _profile_rows(p, lang)

    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    h0 = table.rows[0].cells[0]
    h1 = table.rows[0].cells[1]
    _set_cell_text(h0, s(lang, "profile_metric_header"), bold=True, size=10)
    _set_cell_text(h1, s(lang, "profile_value_header"), bold=True, size=10)
    _shade_cell(h0, HEADER_BG, text_white=True)
    _shade_cell(h1, HEADER_BG, text_white=True)
    for i, (label, value) in enumerate(rows):
        c0 = table.rows[i + 1].cells[0]
        c1 = table.rows[i + 1].cells[1]
        _set_cell_text(c0, label, bold=True, size=10)
        _set_cell_text(c1, value or s(lang, "profile_value_dash"), size=10)
        if i % 2 == 1:
            _shade_cell(c0, ZEBRA_BG)
            _shade_cell(c1, ZEBRA_BG)

    _apply_column_ratios(table, [0.30, 0.70], _usable_width(doc))


def _add_probability_note(doc: Document, profile: StudentProfile, lang: Lang) -> None:
    _h1(doc, s(lang, "note_h1"))

    name = profile.name.split()[0] if profile.name else (
        "the student" if lang == "en" else "학생"
    )
    sat_token = (
        s(lang, "sat_token_with_value", val=profile.sat_total)
        if profile.sat_total
        else s(lang, "sat_token_default")
    )

    intro = doc.add_paragraph()
    intro_run = intro.add_run(s(lang, "note_intro"))
    intro_run.font.size = Pt(10.5)

    bullets = [
        s(lang, "note_bullet_cds"),
        s(lang, "note_bullet_sat_gap", name=name, sat_token=sat_token),
        s(lang, "note_bullet_gpa_rank", name=name),
        s(lang, "note_bullet_premed"),
        s(lang, "note_bullet_qb") if profile.questbridge_candidate else s(lang, "note_bullet_hooks"),
        s(lang, "note_bullet_sffa"),
        s(lang, "note_bullet_oos"),
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    band = doc.add_paragraph()
    band_run = band.add_run(s(lang, "note_band_line"))
    band_run.bold = True
    band_run.font.size = Pt(10.5)
    band.paragraph_format.space_before = Pt(6)
    band.paragraph_format.space_after = Pt(6)

    if profile.sat_total and profile.sat_total < 1500:
        callout = doc.add_paragraph()
        warn = callout.add_run(
            s(lang, "note_warn_sat", sat=profile.sat_total, name=name)
        )
        warn.italic = True
        warn.font.color.rgb = RGBColor.from_string(WARN_TXT)
        warn.font.size = Pt(10.5)
        callout.paragraph_format.left_indent = Inches(0.15)
        callout.paragraph_format.space_before = Pt(8)
        _add_bottom_border(callout, WARN_TXT, sz="6")


# ────────────────────────── Parts 1-3: scope tables ──────────────────────────


def _format_school_name(
    row: CollegeRow,
    *,
    lac_names: frozenset[str] | None = None,
) -> str:
    name = row.name
    if lac_names is not None and _norm_name_key(name) in lac_names:
        name = f"{name} (LAC)"
    if _is_qb(row.name):
        name = f"{name}  ⭐ QB"
    return name


def _norm_name_key(name: str) -> str:
    """Normalise a college name for case/punctuation-insensitive lookup."""
    return re.sub(r"[^a-z0-9]", "", name.lower())


def _add_college_table(
    doc: Document,
    rows: list[CollegeRow],
    lang: Lang,
    *,
    lac_names: frozenset[str] | None = None,
) -> None:
    if not rows:
        _para(doc, s(lang, "tier_no_schools"), italic=True)
        return

    headers = [
        s(lang, "col_num"),
        s(lang, "col_school"),
        s(lang, "col_state"),
        s(lang, "col_odds"),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, h, bold=True, size=10)
        _shade_cell(cell, HEADER_BG, text_white=True)

    sorted_rows = sorted(rows, key=lambda r: r.adjusted_probability)
    for i, row in enumerate(sorted_rows, start=1):
        cells = table.rows[i].cells
        _set_cell_text(cells[0], str(i), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(cells[1], _format_school_name(row, lac_names=lac_names), size=10)
        _set_cell_text(cells[2], row.state or "—", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(
            cells[3],
            f"~{row.adjusted_probability * 100:.0f}%",
            bold=True,
            size=10,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        if i % 2 == 0:
            for c in cells:
                _shade_cell(c, ZEBRA_BG)

    _apply_column_ratios(table, [0.07, 0.65, 0.13, 0.15], _usable_width(doc))


def _add_tier(
    doc: Document, tier: str, n_or_state: str, rows: list[CollegeRow], lang: Lang,
    *,
    use_state_label: bool = False,
    lac_names: frozenset[str] | None = None,
) -> None:
    icon = TIER_ICON[tier]
    color = TIER_COLOR[tier]
    if use_state_label:
        title = s(lang, "tier_heading_state", icon=icon, tier=_tier_label(lang, tier), state=n_or_state)
    else:
        title = s(lang, "tier_heading_n_schools", icon=icon, tier=_tier_label(lang, tier), n=n_or_state)
    _h2(doc, title, color_hex=color)
    _add_college_table(doc, rows, lang, lac_names=lac_names)


def _add_instate_section(
    doc: Document,
    tlist: TieredList,
    profile: StudentProfile,
    lang: Lang,
    *,
    lac_names: frozenset[str] | None = None,
) -> None:
    """Home-state colleges including LACs. Now the FIRST section (Part 1)."""
    state = profile.state or ("Home State" if lang == "en" else "거주 주")
    _h1(doc, s(lang, "part1_h1", state=state))
    _para(doc, s(lang, "part1_intro", state=state))
    _add_tier(doc, "reach", state, tlist.reach, lang, use_state_label=True, lac_names=lac_names)
    _add_tier(doc, "match", state, tlist.match, lang, use_state_label=True, lac_names=lac_names)
    _add_tier(doc, "safety", state, tlist.safety, lang, use_state_label=True, lac_names=lac_names)


def _add_national_section(doc: Document, tlist: TieredList, lang: Lang) -> None:
    """National universities outside the home state. Now the SECOND section (Part 2)."""
    _h1(doc, s(lang, "part2_h1"))
    _add_tier(doc, "reach", str(len(tlist.reach)), tlist.reach, lang)
    _add_tier(doc, "match", str(len(tlist.match)), tlist.match, lang)
    _add_tier(doc, "safety", str(len(tlist.safety)), tlist.safety, lang)


def _add_lac_section(doc: Document, tlist: TieredList, lang: Lang) -> None:
    """LACs nationwide excluding home-state LACs. Now the THIRD section (Part 3)."""
    _h1(doc, s(lang, "part3_h1"))
    _para(doc, s(lang, "part3_intro"))
    _add_tier(doc, "reach", str(len(tlist.reach)), tlist.reach, lang)
    _add_tier(doc, "match", str(len(tlist.match)), tlist.match, lang)
    _add_tier(doc, "safety", str(len(tlist.safety)), tlist.safety, lang)


# ────────────────────────── Parts 4-5: ED / EA ──────────────────────────


def _round_label_for_ed(row: CollegeRow, lang: Lang) -> str:
    return s(lang, "round_ed")


def _round_label_for_ea(row: CollegeRow, lang: Lang) -> str:
    if row.has_rea_or_scea:
        return s(lang, "round_rea_scea")
    return s(lang, "round_ea")


def _add_early_table(
    doc: Document,
    rows: list[CollegeRow],
    round_label_col_key: str,
    round_label_fn,
    lang: Lang,
) -> None:
    if not rows:
        _para(doc, s(lang, "tier_no_early"), italic=True)
        return

    headers = [
        s(lang, "col_school"),
        s(lang, "col_state"),
        s(lang, round_label_col_key),
        s(lang, "col_odds"),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, h, bold=True, size=10)
        _shade_cell(cell, HEADER_BG, text_white=True)

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        _set_cell_text(cells[0], _format_school_name(row), size=10)
        _set_cell_text(cells[1], row.state or "—", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(cells[2], round_label_fn(row, lang), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(
            cells[3],
            f"~{row.adjusted_probability * 100:.0f}%",
            bold=True,
            size=10,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        if i % 2 == 0:
            for c in cells:
                _shade_cell(c, ZEBRA_BG)

    _apply_column_ratios(table, [0.55, 0.13, 0.17, 0.15], _usable_width(doc))


def _add_part4_ed(doc: Document, ed_rows: list[CollegeRow], lang: Lang) -> None:
    _h1(doc, s(lang, "part4_h1"))
    _para(doc, s(lang, "part4_intro"))
    top = sorted(ed_rows, key=lambda r: -r.adjusted_probability)[:ED_MAX]
    _add_early_table(doc, top, "col_ed_round", _round_label_for_ed, lang)


def _add_part5_ea(doc: Document, ea_rows: list[CollegeRow], lang: Lang) -> None:
    _h1(doc, s(lang, "part5_h1"))
    _para(doc, s(lang, "part5_intro"))
    top = sorted(ea_rows, key=lambda r: -r.adjusted_probability)[:EA_MAX]
    _add_early_table(doc, top, "col_ea_type", _round_label_for_ea, lang)


# ────────────────────────── Part 6: Action plan ──────────────────────────


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _emit_inline_runs(paragraph, text: str) -> None:
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        bold_run = paragraph.add_run(m.group(1))
        bold_run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _render_action_markdown(doc: Document, md_text: str) -> None:
    lines = md_text.splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            _h3(doc, line[4:].strip())
        elif line.startswith("## "):
            _h2(doc, line[3:].strip())
        elif line.startswith("# "):
            _h2(doc, line[2:].strip())
        elif line.lstrip().startswith(("- ", "* ")):
            bullet_text = line.lstrip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            _emit_inline_runs(p, bullet_text)
        elif re.match(r"^\d+\.\s", line.lstrip()):
            num_text = re.sub(r"^\d+\.\s", "", line.lstrip())
            p = doc.add_paragraph(style="List Number")
            _emit_inline_runs(p, num_text)
        else:
            p = doc.add_paragraph()
            _emit_inline_runs(p, line)


def _add_part6_action_plan(doc: Document, markdown: str, lang: Lang) -> None:
    _h1(doc, s(lang, "part6_h1"))
    _render_action_markdown(doc, markdown)


# ────────────────────────── Glossary + Closing ──────────────────────────


def _add_glossary(doc: Document, lang: Lang) -> None:
    _h1(doc, s(lang, "glossary_h1"))
    glossary = get_glossary(lang)
    table = doc.add_table(rows=1 + len(glossary), cols=2)
    headers = [s(lang, "glossary_term_header"), s(lang, "glossary_definition_header")]
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        _set_cell_text(c, h, bold=True, size=10)
        _shade_cell(c, HEADER_BG, text_white=True)
    for i, (term, definition) in enumerate(glossary, start=1):
        cells = table.rows[i].cells
        _set_cell_text(cells[0], term, bold=True, size=10)
        _set_cell_text(cells[1], definition, size=10)
        if i % 2 == 0:
            for c in cells:
                _shade_cell(c, ZEBRA_BG)
    _apply_column_ratios(table, [0.18, 0.82], _usable_width(doc))


def _add_closing_line(doc: Document, profile: StudentProfile, lang: Lang) -> None:
    track = profile.career_goal or profile.intended_major or s(lang, "closing_default_track")
    today = _dt.date.today()
    if today.month >= 7:
        y0, y1 = today.year, (today.year + 1) % 100
    else:
        y0, y1 = today.year - 1, today.year % 100
    ay = s(lang, "ay_template", y0=y0, y1=f"{y1:02d}")

    p = doc.add_paragraph()
    run = p.add_run(s(lang, "closing_line", track=track, name=profile.name, ay=ay))
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(SUBTITLE_GRAY)
    p.paragraph_format.space_before = Pt(20)


# ────────────────────────── public API ──────────────────────────


def build_docx(
    *,
    profile: StudentProfile,
    national: TieredList,
    instate: TieredList,
    lac: TieredList,
    ed_rows: list[CollegeRow],
    ea_rows: list[CollegeRow],
    action_plan_md: str,
    flags: list[ValidationFlag],
    out_path: Path,
    lang: Lang = "en",
    lac_names: frozenset[str] | None = None,
) -> None:
    """Build the Word report.

    `lac_names` should be a set of normalised college-name keys (lowercased
    alphanumerics) for every Liberal Arts College known to the corpus. Used
    only in Part 3 (home-state list) to tag in-state LACs with "(LAC)".
    """
    doc = Document()
    _setup_document(doc, profile, lang)

    # Page 1: Title + Profile + Note
    _add_title_block(doc, profile, lang)
    _add_profile_snapshot(doc, profile, lang)
    _add_probability_note(doc, profile, lang)

    _page_break(doc)
    _add_instate_section(doc, instate, profile, lang, lac_names=lac_names)

    _page_break(doc)
    _add_national_section(doc, national, lang)

    _page_break(doc)
    _add_lac_section(doc, lac, lang)

    _page_break(doc)
    _add_part4_ed(doc, ed_rows, lang)

    _page_break(doc)
    _add_part5_ea(doc, ea_rows, lang)

    _page_break(doc)
    _add_part6_action_plan(doc, action_plan_md, lang)

    # Validation notes intentionally not rendered (debug-only signal kept in console).

    _page_break(doc)
    _add_glossary(doc, lang)
    _add_closing_line(doc, profile, lang)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
